"""Standalone DOCX compliance report generator.

Uses python-docx (optional dependency) to generate professional compliance
reports without any Kuzu/monorepo dependency.

Install: pip install bespoketracker-comply[docx]
"""
from __future__ import annotations

import io
import logging
from datetime import datetime, timezone
from typing import Optional

log = logging.getLogger(__name__)


def generate_comply_docx(report: dict) -> Optional[bytes]:
    """Generate a DOCX compliance report from a scan report dict.

    Parameters
    ----------
    report : dict
        Comply scan report (from run_comply_scan or stored scan).

    Returns
    -------
    bytes or None
        DOCX file content, or None if python-docx is not installed.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        log.info("python-docx not installed; DOCX export unavailable")
        return None

    doc = Document()

    # ── Styles ────────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # ── Title page ────────────────────────────────────────────────
    framework = report.get("framework", "unknown")
    target = report.get("target", "")
    summary = report.get("summary", {})
    score = summary.get("score", 0)
    generated = report.get("generatedAt", datetime.now(timezone.utc).isoformat())

    title = doc.add_heading("Compliance Assessment Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    meta = doc.add_table(rows=5, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell(meta, 0, 0, "Framework:", bold=True)
    _set_cell(meta, 0, 1, _framework_label(framework))
    _set_cell(meta, 1, 0, "Target:", bold=True)
    _set_cell(meta, 1, 1, target)
    _set_cell(meta, 2, 0, "Scan Depth:", bold=True)
    _set_cell(meta, 2, 1, report.get("scanDepth", "content"))
    _set_cell(meta, 3, 0, "Score:", bold=True)
    _set_cell(meta, 3, 1, f"{score:.1f}%")
    _set_cell(meta, 4, 0, "Generated:", bold=True)
    _set_cell(meta, 4, 1, generated[:19])

    doc.add_page_break()

    # ── Executive Summary ─────────────────────────────────────────
    doc.add_heading("1. Executive Summary", level=1)

    met = summary.get("met", 0)
    partial = summary.get("partial", 0)
    gap = summary.get("gap", 0)
    total = summary.get("total", 0)

    doc.add_paragraph(
        f"This report presents the compliance assessment of the codebase at "
        f"{target} against the {_framework_label(framework)} framework. "
        f"The scan analysed {report.get('model', {}).get('filesScanned', 0)} files "
        f"and discovered {report.get('model', {}).get('featuresDiscovered', 0)} features."
    )

    doc.add_paragraph(
        f"Of {total} controls evaluated, {met} are fully met, "
        f"{partial} are partially met, and {gap} have gaps requiring attention. "
        f"The overall compliance score is {score:.1f}%."
    )

    # Summary table
    summary_table = doc.add_table(rows=2, cols=4)
    summary_table.style = 'Light List Accent 1'
    _set_cell(summary_table, 0, 0, "Met", bold=True)
    _set_cell(summary_table, 0, 1, "Partial", bold=True)
    _set_cell(summary_table, 0, 2, "Gap", bold=True)
    _set_cell(summary_table, 0, 3, "Total", bold=True)
    _set_cell(summary_table, 1, 0, str(met))
    _set_cell(summary_table, 1, 1, str(partial))
    _set_cell(summary_table, 1, 2, str(gap))
    _set_cell(summary_table, 1, 3, str(total))

    # ── Per-Article Assessment ────────────────────────────────────
    doc.add_heading("2. Detailed Assessment", level=1)

    for i, article in enumerate(report.get("articles", []), 1):
        art_id = article.get("articleId", "")
        art_label = article.get("label", "")
        doc.add_heading(f"2.{i}. {art_label} (Art. {art_id})", level=2)

        art_met = article.get("met", 0)
        art_partial = article.get("partial", 0)
        art_gap = article.get("gap", 0)

        doc.add_paragraph(
            f"Controls: {art_met} met, {art_partial} partial, {art_gap} gap"
        )

        # Controls table
        controls = article.get("controls", [])
        if controls:
            tbl = doc.add_table(rows=len(controls) + 1, cols=4)
            tbl.style = 'Light List Accent 1'
            _set_cell(tbl, 0, 0, "Control", bold=True)
            _set_cell(tbl, 0, 1, "Requirement", bold=True)
            _set_cell(tbl, 0, 2, "Status", bold=True)
            _set_cell(tbl, 0, 3, "Evidence / Recommendation", bold=True)

            for j, ctrl in enumerate(controls, 1):
                _set_cell(tbl, j, 0, ctrl.get("controlId", ""))
                req = ctrl.get("requirement", "")
                if len(req) > 200:
                    req = req[:197] + "..."
                _set_cell(tbl, j, 1, req)

                status = ctrl.get("status", "gap")
                _set_cell(tbl, j, 2, status.upper(),
                          color=_status_color(status))

                evidence = ctrl.get("evidence", [])
                rec = ctrl.get("recommendation", "")
                details = "\n".join(evidence) if evidence else ""
                if rec:
                    details += f"\nRecommendation: {rec}" if details else f"Recommendation: {rec}"
                _set_cell(tbl, j, 3, details)

    # ── Remediation Recommendations ───────────────────────────────
    gap_controls = []
    for article in report.get("articles", []):
        for ctrl in article.get("controls", []):
            if ctrl.get("status") == "gap":
                gap_controls.append(ctrl)

    if gap_controls:
        doc.add_heading("3. Remediation Recommendations", level=1)
        doc.add_paragraph(
            f"The following {len(gap_controls)} controls have gaps that require attention:"
        )
        for ctrl in gap_controls:
            ctrl_id = ctrl.get("controlId", "")
            rec = ctrl.get("recommendation", "No specific recommendation")
            p = doc.add_paragraph(f"{ctrl_id}: ", style='List Number')
            p.add_run(rec)
    else:
        doc.add_heading("3. Remediation Recommendations", level=1)
        doc.add_paragraph("All controls are fully or partially met. No critical gaps identified.")

    # ── Footer ────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("4. About This Report", level=1)
    doc.add_paragraph(
        "This report was generated by BespokeTracker Comply, an automated "
        "compliance gap analysis tool. The assessment is based on static code "
        "analysis and pattern matching against regulatory framework controls."
    )
    doc.add_paragraph(
        "This report should be reviewed by a qualified compliance professional. "
        "Automated assessments provide evidence indicators but cannot replace "
        "human judgement for regulatory compliance determinations."
    )
    doc.add_paragraph(f"Report generated: {generated[:19]}")
    doc.add_paragraph(f"Scan ID: {report.get('scanId', '')}")
    doc.add_paragraph(f"Commit: {report.get('commitSha', 'N/A')}")

    # ── Write to bytes ────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _set_cell(table, row: int, col: int, text: str,
              bold: bool = False, color: Optional[tuple] = None):
    """Set a table cell's text with optional formatting."""
    cell = table.cell(row, col)
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = bold
            if color:
                from docx.shared import RGBColor
                run.font.color.rgb = RGBColor(*color)


def _status_color(status: str) -> Optional[tuple]:
    return {
        "met": (63, 185, 80),
        "partial": (210, 153, 34),
        "gap": (248, 81, 73),
    }.get(status)


_FRAMEWORK_LABELS = {
    "eu-ai-act": "EU AI Act (Regulation 2024/1689)",
    "eu_ai_act": "EU AI Act (Regulation 2024/1689)",
    "nist-ai-rmf": "NIST AI Risk Management Framework 1.0",
    "nist_ai_rmf": "NIST AI Risk Management Framework 1.0",
    "iso-42001": "ISO/IEC 42001:2023",
    "iso_42001": "ISO/IEC 42001:2023",
    "colorado-sb-24-205": "Colorado SB 24-205 (Consumer Protections)",
    "colorado_sb_24_205": "Colorado SB 24-205 (Consumer Protections)",
    "soc2-ai": "SOC 2 (AI Trust Services Criteria)",
    "soc2_ai": "SOC 2 (AI Trust Services Criteria)",
}


def _framework_label(fw: str) -> str:
    return _FRAMEWORK_LABELS.get(fw, fw)
